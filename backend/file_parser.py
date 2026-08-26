import io
import os
import re
import shutil
import tempfile
from typing import Optional
from PIL import Image, ImageEnhance, ImageFilter, ImageOps

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "image/png",
    "image/jpeg",
}
MAX_FILE_SIZE = 10 * 1024 * 1024
MAX_OCR_PAGES = 15


def _configure_tesseract() -> bool:
    try:
        import pytesseract

        # 1. Check environment variable
        env_path = os.getenv("TESSERACT_PATH", "").strip()
        if env_path and os.path.exists(env_path):
            pytesseract.pytesseract.tesseract_cmd = env_path
            return True

        # 2. Check system PATH
        which_path = shutil.which("tesseract")
        if which_path:
            pytesseract.pytesseract.tesseract_cmd = which_path
            return True

        # 3. Check standard Windows installation locations
        windows_candidates = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            os.path.expandvars(r"%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe"),
            os.path.expandvars(r"%USERPROFILE%\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"),
        ]
        for cand in windows_candidates:
            if cand and os.path.exists(cand):
                pytesseract.pytesseract.tesseract_cmd = cand
                return True

        return False
    except Exception:
        return False


def _clean_extracted_text(text: str) -> str:
    if not text:
        return ""
    # Strip non-printable and replacement characters
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f\ufffd]", "", text)
    # Fix broken hyphenations across newlines (e.g., 'analy-\nsis' -> 'analysis')
    text = re.sub(r"(\b[A-Za-z]+)-\s*\n\s*([A-Za-z]+\b)", r"\1\2", text)
    # Remove lines that contain only noisy stray punctuation artifacts
    cleaned_lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
        elif re.search(r"[A-Za-z0-9]", stripped):
            cleaned_lines.append(stripped)
    text = "\n".join(cleaned_lines)
    # Normalize excessive consecutive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _score_ocr_text(text: str) -> float:
    if not text:
        return 0.0
    words = re.findall(r"\b[A-Za-z0-9$%.,-]{2,}\b", text)
    if not words:
        return 0.0
    total_chars = len(text)
    alpha_num_count = sum(c.isalnum() for c in text)
    ratio = alpha_num_count / total_chars if total_chars else 0
    return len(words) * (ratio ** 1.2)


def _otsu_threshold(gray_img: Image.Image) -> Image.Image:
    ac = ImageOps.autocontrast(gray_img, cutoff=(1, 1))
    hist = ac.histogram()
    total = sum(hist)
    current_max, threshold = 0.0, 128
    total_mean = sum(i * hist[i] for i in range(256))
    sum_back, weight_back = 0.0, 0
    for i in range(256):
        weight_back += hist[i]
        if weight_back == 0:
            continue
        weight_fore = total - weight_back
        if weight_fore == 0:
            break
        sum_back += i * hist[i]
        mean_back = sum_back / weight_back
        mean_fore = (total_mean - sum_back) / weight_fore
        var_between = weight_back * weight_fore * ((mean_back - mean_fore) ** 2)
        if var_between > current_max:
            current_max = var_between
            threshold = i
    bw = ac.point(lambda p: 255 if p > threshold else 0, mode="1")
    return bw.convert("L")


def _ocr_pil_image(img: Image.Image) -> str:
    try:
        import pytesseract

        _configure_tesseract()

        # 1. Add white border padding to prevent edge text clipping by Tesseract
        img = ImageOps.expand(img, border=30, fill="white")

        # 2. Upscale if image resolution is low (ensures optimal ~200 DPI for Tesseract)
        w, h = img.size
        max_dim = max(w, h)
        min_dim = min(w, h)
        if max_dim < 1400 or min_dim < 900:
            scale = max(1.5, 1600.0 / max_dim)
            img = img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)

        # 3. Convert to grayscale, auto-contrast, and sharpen
        gray = ImageOps.grayscale(img)
        contrasted = ImageOps.autocontrast(gray, cutoff=(1, 1))
        enhanced = ImageEnhance.Contrast(contrasted).enhance(1.6)
        sharpened = enhanced.filter(ImageFilter.UnsharpMask(radius=1.5, percent=150, threshold=2))
        binarized = _otsu_threshold(sharpened)

        # 4. Detect dark backgrounds (inverted text)
        mean_lum = sum(gray.histogram()[i] * i for i in range(256)) / (gray.size[0] * gray.size[1])
        variants = [sharpened, binarized]
        if mean_lum < 110:
            inverted = ImageOps.invert(sharpened)
            variants.extend([inverted, _otsu_threshold(inverted)])

        # 5. Multi-pass candidate scoring across PSM modes
        candidates = []
        for variant in variants:
            for psm in ["3", "6"]:
                try:
                    res = pytesseract.image_to_string(variant, config=f"--oem 3 --psm {psm}")
                    cleaned = _clean_extracted_text(res)
                    if cleaned:
                        score = _score_ocr_text(cleaned)
                        candidates.append((score, cleaned))
                except Exception:
                    continue

        if not candidates:
            return ""

        candidates.sort(key=lambda x: x[0], reverse=True)
        return candidates[0][1]
    except Exception:
        return ""


def _allowed(filename: str) -> bool:
    _, ext = os.path.splitext(filename)
    return ext.lower() in ALLOWED_EXTENSIONS


def _allowed_mime(mime: str) -> bool:
    return mime.lower() in ALLOWED_MIME_TYPES


def _read_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def _read_pdf(file_bytes: bytes) -> str:
    extracted_pages = []
    pdf_doc = None

    try:
        import pypdfium2 as pdfium
        pdf_doc = pdfium.PdfDocument(file_bytes)
    except Exception:
        pdf_doc = None

    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            total_pages = len(pdf.pages)
            ocr_page_count = 0
            for i, page in enumerate(pdf.pages):
                direct_text = _clean_extracted_text(page.extract_text() or "")
                has_images = bool(page.images)

                alpha_count = sum(c.isalnum() for c in direct_text)
                needs_ocr = False
                if alpha_count < 25 and pdf_doc is not None and ocr_page_count < MAX_OCR_PAGES:
                    needs_ocr = True
                elif has_images and pdf_doc is not None and ocr_page_count < MAX_OCR_PAGES:
                    needs_ocr = True

                if needs_ocr:
                    try:
                        pil_img = pdf_doc[i].render(scale=2.0).to_pil()
                        ocr_text = _ocr_pil_image(pil_img)
                        ocr_page_count += 1
                        if alpha_count < 25:
                            if ocr_text:
                                direct_text = f"{direct_text}\n{ocr_text}".strip() if direct_text else ocr_text
                        elif len(ocr_text) > len(direct_text) + 25:
                            direct_text = ocr_text
                    except Exception:
                        pass

                if direct_text:
                    if total_pages > 1:
                        extracted_pages.append(f"--- Page {i + 1} ---\n{direct_text}")
                    else:
                        extracted_pages.append(direct_text)

    except Exception:
        # Fallback if pdfplumber fails: use pypdfium2 + high-res OCR exclusively
        if pdf_doc is not None:
            total_pages = len(pdf_doc)
            ocr_page_count = 0
            for i in range(total_pages):
                if ocr_page_count >= MAX_OCR_PAGES:
                    break
                try:
                    pil_img = pdf_doc[i].render(scale=2.0).to_pil()
                    ocr_text = _ocr_pil_image(pil_img)
                    ocr_page_count += 1
                    if ocr_text:
                        if total_pages > 1:
                            extracted_pages.append(f"--- Page {i + 1} ---\n{ocr_text}")
                        else:
                            extracted_pages.append(ocr_text)
                except Exception:
                    pass
    final_text = "\n\n".join(p for p in extracted_pages if p).strip()
    if not final_text:
        # Last attempt: render all pages with pypdfium2 if not already done
        if pdf_doc is not None:
            total_pages = len(pdf_doc)
            ocr_page_count = 0
            for i in range(total_pages):
                if ocr_page_count >= MAX_OCR_PAGES:
                    break
                try:
                    pil_img = pdf_doc[i].render(scale=2.0).to_pil()
                    ocr_text = _ocr_pil_image(pil_img)
                    ocr_page_count += 1
                    if ocr_text:
                        if total_pages > 1:
                            extracted_pages.append(f"--- Page {i + 1} ---\n{ocr_text}")
                        else:
                            extracted_pages.append(ocr_text)
                except Exception:
                    pass
            final_text = "\n\n".join(p for p in extracted_pages if p).strip()

    if not final_text:
        raise ValueError("No readable text could be extracted or OCR-scanned from this PDF.")
    return _clean_extracted_text(final_text)


def _read_docx(file_bytes: bytes) -> str:
    try:
        import docx

        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp.flush()
                tmp_path = tmp.name
            doc = docx.Document(tmp_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            # Extract text from tables in docx
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return _clean_extracted_text("\n".join(paragraphs))
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {e}")


def _read_image(file_bytes: bytes) -> str:
    try:
        _configure_tesseract()
        image = Image.open(io.BytesIO(file_bytes))
        text = _ocr_pil_image(image)
        if not text:
            raise ValueError("No text could be recognized in the uploaded image.")
        return _clean_extracted_text(text)
    except Exception as e:
        raise RuntimeError(f"Failed to read image (OCR): {e}")


def extract_text(filename: str, content: bytes, mime: str = "") -> str:
    if len(content) > MAX_FILE_SIZE:
        raise ValueError("File size exceeds 10MB limit")
    if not _allowed(filename):
        raise ValueError(f"Unsupported file type. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}")
    if mime and not _allowed_mime(mime):
        raise ValueError(f"Unsupported MIME type: {mime}. Allowed: {', '.join(sorted(ALLOWED_MIME_TYPES))}")

    _, ext = os.path.splitext(filename)
    ext = ext.lower()

    if ext == ".txt":
        return _read_txt(content)
    if ext == ".pdf":
        return _read_pdf(content)
    if ext == ".docx":
        return _read_docx(content)
    if ext in (".png", ".jpg", ".jpeg"):
        return _read_image(content)

    raise ValueError(f"Unsupported file type: {ext}")
